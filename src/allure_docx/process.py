import os
from configparser import ConfigParser
from os import listdir
from os.path import join, isfile
from time import ctime
from datetime import timedelta, datetime
import warnings
import re

import json

from docx.shared import Mm, Cm
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from . import piechart

INDENT = 6


def _format_argval(argval):
    """Remove newlines and limit max length

    From Allure-pytest logger (formats argument in the CLI live logs).
    Consider using the same function."""

    MAX_ARG_LENGTH = 100
    argval = argval.replace("\n", " ")
    if len(argval) > MAX_ARG_LENGTH:
        argval = argval[:3] + " ... " + argval[-MAX_ARG_LENGTH:]
    return argval


def get_config_from_file(standard_path, path):
    def transform_by_status_to_dict(config, section):
        section_old = config[section]
        config[section] = {}
        config[section]["failed"] = []
        config[section]["broken"] = []
        config[section]["passed"] = []
        config[section]["skipped"] = []
        config[section]["unknown"] = []
        for key in section_old:
            if "f" in section_old[key]:
                config[section]["failed"].append(key)
            if "b" in section_old[key]:
                config[section]["broken"].append(key)
            if "p" in section_old[key]:
                config[section]["passed"].append(key)
            if "s" in section_old[key]:
                config[section]["skipped"].append(key)
            if "u" in section_old[key]:
                config[section]["unknown"].append(key)

    try:
        config_parser = ConfigParser()
        config_parser.read(standard_path)
        if path is not standard_path:
            config_parser.read(path)
        config = {s: dict(config_parser.items(s)) for s in config_parser.sections()}
        transform_by_status_to_dict(config, "info")
        transform_by_status_to_dict(config, "labels")
        return config
    except Exception as e:
        print("Failed to read config file: " + str(e))
        exit(-1)


def build_data(alluredir):
    def _process_steps(session, node):
        if "start" in node:
            if "start" not in session:
                session["start"] = node["start"]
            elif session["start"] is None:
                session["start"] = node["start"]
            elif node["start"] < session["start"]:
                session["start"] = node["start"]

        if "stop" in node:
            if "stop" not in session:
                session["stop"] = node["stop"]
            elif session["stop"] is None:
                session["stop"] = node["stop"]
            elif node["stop"] > session["stop"]:
                session["stop"] = node["stop"]

        if "steps" in node:
            for step in node["steps"]:
                _process_steps(session, step)

    json_results = [f for f in listdir(alluredir) if isfile(join(alluredir, f)) and "result" in f]
    json_containers = [f for f in listdir(alluredir) if isfile(join(alluredir, f)) and "container" in f]

    session = {
        "alluredir": alluredir,
        "start": None,
        "stop": None,
        "results": {
            "broken": 0,
            "failed": 0,
            "skipped": 0,
            "passed": 0,
        },
        "results_relative": {
            "broken": 0,
            "failed": 0,
            "skipped": 0,
            "passed": 0,
        },
        "total": 0,
    }

    data_containers = []
    for file in json_containers:
        with open(join(alluredir, file), encoding="utf-8") as f:
            container = json.load(f)
            data_containers.append(container)

    data_results = []
    for file in json_results:
        with open(join(alluredir, file), encoding="utf-8") as f:
            result = json.load(f)
            result["_lastmodified"] = os.path.getmtime(join(alluredir, file))

            skip = False
            for previous_item in list(data_results):  # copy
                if previous_item["name"] == result["name"]:
                    if previous_item["_lastmodified"] > result["_lastmodified"]:
                        skip = True
                    else:
                        data_results.remove(previous_item)
                    break
            if skip:
                continue
            data_results.append(result)

    for result in data_results:
        _process_steps(session, result)
        session["total"] += 1
        session["results"][result["status"]] += 1

        result["parents"] = []
        for container in data_containers:
            if "children" not in container:
                continue
            if result["uuid"] in container["children"]:
                result["parents"].append(container)
                if "befores" in container:
                    for before in container["befores"]:
                        _process_steps(session, before)
                if "afters" in container:
                    for after in container["afters"]:
                        _process_steps(session, after)

    if session["total"] == 0:
        warnings.warn("No test result files were found!")

    def getsortingkey(d):
        classification = {"broken": 0, "failed": 1, "skipped": 2, "passed": 3}
        return "{}-{}".format(classification[d["status"]], d["name"])

    sorted_results = sorted(data_results, key=getsortingkey)

    if session["start"] is not None:
        session["duration"] = str(timedelta(seconds=(session["stop"] - session["start"]) / 1000.0))
        session["start"] = ctime(session["start"] / 1000.0)
        session["stop"] = ctime(session["stop"] / 1000.0)
    else:
        session["duration"] = "Not available"
        session["start"] = "Not available"
        session["stop"] = "Not available"

    for item in session["results"]:
        if session["total"] > 0:
            session["results_relative"][item] = "{:.2f}%".format(100 * session["results"][item] / session["total"])
        else:
            session["results_relative"][item] = "Not available"

    return sorted_results, session

def create_docx(sorted_results, session, config):
    def create_TOC(document):
        # Snippet from:
        # https://github.com/python-openxml/python-docx/issues/36
        paragraph = document.add_paragraph()
        run = paragraph.add_run()
        fldChar = OxmlElement("w:fldChar")  # creates a new element
        fldChar.set(qn("w:fldCharType"), "begin")  # sets attribute on element
        instrText = OxmlElement("w:instrText")
        instrText.set(qn("xml:space"), "preserve")  # sets attribute on element
        instrText.text = 'TOC \\o "1-1" \\h \\z'  # change 1-3 depending on heading levels you need

        fldChar2 = OxmlElement("w:fldChar")
        fldChar2.set(qn("w:fldCharType"), "separate")
        fldChar3 = OxmlElement("w:t")
        fldChar3.text = "Right-click to update field."
        fldChar2.append(fldChar3)

        fldChar4 = OxmlElement("w:fldChar")
        fldChar4.set(qn("w:fldCharType"), "end")

        r_element = run._r
        r_element.append(fldChar)
        r_element.append(instrText)
        r_element.append(fldChar2)
        r_element.append(fldChar4)
        p_element = paragraph._p

    def print_attachments(document, item):
        if "attachments" in item:
            for attachment in item["attachments"]:
                document.add_paragraph("[Attachment] {}".format(attachment["name"]), style="Step")
                if "image" in attachment["type"]:
                    document.add_picture(
                        os.path.join(session["alluredir"], attachment["source"]),
                        width=Mm(100),
                    )
                    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    def print_steps(document, parent_step, config_info, indent=0):
        indent_str = indent * INDENT * " "
        if "steps" in parent_step:
            for step in parent_step["steps"]:
                if step["status"] in ["failed", "broken"]:
                    stepstyle = "Step Failed"
                else:
                    stepstyle = "Step"
                document.add_paragraph("{}> {}".format(indent_str, step["name"]), style=stepstyle)
                if "parameters" in config_info and "parameters" in step:
                    for p in step["parameters"]:
                        paragraph = document.add_paragraph("{}    ".format(indent_str), style="Step Param Parag")
                        paragraph.add_run(
                            "{} = {}".format(p["name"], _format_argval(p["value"])),
                            style="Step Param",
                        )
                if (
                        "details" in config_info
                        and "statusDetails" in step
                        and len(step["statusDetails"]) != 0
                        and (
                        "message" in step["statusDetails"]
                        and len(step["statusDetails"]["message"]) != 0
                        or "trace" in config_info
                        and "trace" in step["statusDetails"]
                        and len(step["statusDetails"]["trace"]) != 0
                )
                ):
                    if "message" in step["statusDetails"]:
                        document.add_paragraph(step["statusDetails"]["message"], style=stepstyle)
                    if "trace" in config_info and "trace" in step["statusDetails"]:
                        table = document.add_table(rows=1, cols=1, style="Trace table")
                        hdr_cells = table.rows[0].cells
                        hdr_cells[0].add_paragraph(step["statusDetails"]["trace"] + "\n", style="Code")
                if "attachments" in config_info:
                    print_attachments(document, step)
                print_steps(document, step, config_info, indent + 1)

    def delete_paragraph(paragraph):
        p = paragraph._element
        p.getparent().remove(p)
        p._p = p._element = None

    def add_field(run, field):
        def create_attribute(element, name, value):
            element.set(qn(name), value)

        def create_element(name):
            return OxmlElement(name)

        fldChar1 = create_element('w:fldChar')
        create_attribute(fldChar1, 'w:fldCharType', 'begin')

        instrText = create_element('w:instrText')
        create_attribute(instrText, 'xml:space', 'preserve')
        instrText.text = field

        fldChar2 = create_element('w:fldChar')
        create_attribute(fldChar2, 'w:fldCharType', 'end')

        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)

    def create_footer(footer):
        footer.paragraphs[0].text += datetime.today().strftime('%Y-%m-%d')
        footer.paragraphs[0].text += "\t\t"
        footer_run = footer.paragraphs[0].add_run()
        add_field(footer_run, field="PAGE")

    def create_header(header, details=False):
        htable = header.add_table(1, 2, Cm(16))
        htable.style = "header table"
        htab_cells = htable.rows[0].cells

        if config['logo']['path'] is not None:
            ht1 = htab_cells[1].add_paragraph()
            ht1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            kh = ht1.add_run()
            kh.add_picture(config['logo']['path'], width=Cm(5))

        if details:
            header_text = "Test Report"
            header_text += "\n" + config['cover']['title']
            if 'device_under_test' in config['details']:
                header_text += "\n" + config['details']['device_under_test']
            htab_cells[0].add_paragraph(header_text)
            delete_paragraph(header.paragraphs[0])
            delete_paragraph(htab_cells[0].paragraphs[0])
            delete_paragraph(htab_cells[1].paragraphs[0])
            header.add_paragraph("")

    document = Document(config['template_path'])

    header = document.sections[0].header
    create_header(header)

    delete_paragraph(document.paragraphs[-0])
    if 'company_name' in config['cover']:
        document.add_paragraph("\n" + config['cover']['company_name'], style="company")
    document.add_paragraph("\n\n\n\nTest Report", style="Title")
    subtitle = config['cover']['title']
    if 'device_under_test' in config['details']:
        subtitle += "\n" + config['details']['device_under_test']
    document.add_paragraph(subtitle, style="Subtitle")
    document.add_paragraph("\n" + datetime.today().strftime('%Y-%m-%d'), style="heading 2")

    document.add_section()
    footer = document.sections[1].footer
    footer.is_linked_to_previous = False
    create_footer(footer)
    header = document.sections[1].header
    header.is_linked_to_previous = False
    create_header(header, True)

    if 'details' in config and len(config['details']) > 0:
        document.add_paragraph("Test Details", style="Heading 1")

        i = 0
        detail_table = document.add_table(rows=len(config['details']), cols=2, style="Label table")
        for detail in config['details'].items():
            detail_table.rows[i].cells[0].paragraphs[-1].clear().add_run(detail[0].replace("_", " ").capitalize())
            detail_table.rows[i].cells[1].paragraphs[-1].clear().add_run(re.sub(r";\s*", "\n", detail[1]))
            i += 1

    document.add_paragraph("Test Session Summary", style="Heading 1")

    if not sorted_results:
        document.add_paragraph("No test result files were found.")
        document.save(config['output_filename'])
        return
    table = document.add_table(rows=1, cols=2)
    summary_cell = table.rows[0].cells[0]
    summary_cell.add_paragraph(
        "Start: {}\nEnd: {}\nDuration: {}".format(session["start"], session["stop"], session["duration"])
    )
    delete_paragraph(summary_cell.paragraphs[0])

    results_strs = []
    for item in session["results"]:
        results_strs.append("{}: {} ({})".format(item, session["results"][item], session["results_relative"][item]))
    summary_cell.add_paragraph("\n".join(results_strs))

    piechart_cell = table.rows[0].cells[1]
    paragraph = piechart_cell.paragraphs[0]
    run = paragraph.add_run()
    run.add_picture(session["piechart_source"], width=Mm(75))

    document.add_page_break()

    # print tests
    for test in sorted_results:
        # config elements for the specific status of this test
        config_info = config["info"][test["status"]]
        config_labels = config["labels"][test["status"]]

        document.add_paragraph(
            "{}  [ {} ]".format(test["name"], test["status"]),
            style="Heading {}".format(test["status"]),
        )

        table = None
        added_table = False
        if "duration" in config_info:
            duration = test["stop"] - test["start"]
            duration_unit = "ms"
            if duration > 1000:
                duration_unit = "s"
                duration = duration / 1000
                if duration > 60:
                    duration_unit = "min"
                    duration = duration / 60

            table = document.add_table(rows=1, cols=2, style="Label table")
            table.rows[0].cells[0].paragraphs[-1].clear().add_run("Duration")
            table.rows[0].cells[1].paragraphs[-1].clear().add_run(str(duration) + duration_unit)
            added_table = True

        # add labels to table
        for label_name in config_labels:
            if not added_table:
                table = document.add_table(rows=0, cols=2, style="Label table")
                added_table = True
            iterator = iter(label for label in test["labels"] if label["name"].lower() == label_name)
            label = next(iterator, None)
            if label is not None:
                row = table.add_row()
                row.cells[0].paragraphs[-1].clear().add_run(label_name.capitalize())
                while label is not None:
                    row.cells[1].add_paragraph(label["value"])
                    label = next(iterator, None)
                delete_paragraph(row.cells[1].paragraphs[0])

        if table is not None:
            table.columns[0].width = Cm(4)
            for cell in table.columns[0].cells:
                cell.width = Cm(4)
            table.columns[1].width = Cm(12)
            for cell in table.columns[1].cells:
                cell.width = Cm(12)
            document.add_paragraph()

        if "description" in config_info:
            document.add_heading("Description", level=2)
            if "description" in test and len(test["description"]) != 0:
                document.add_paragraph(test["description"])
            else:
                document.add_paragraph("No description available.")

        if "parameters" in config_info and "parameters" in test and len(test["parameters"]) != 0:
            document.add_heading("Parameters", level=2)
            for p in test["parameters"]:
                document.add_paragraph("{}: {}".format(p["name"], p["value"]), style="Step")

        if (
                "details" in config_info
                and "statusDetails" in test
                and len(test["statusDetails"]) != 0
                and (
                "message" in test["statusDetails"]
                and len(test["statusDetails"]["message"]) != 0
                or "trace" in config_info
                and "trace" in test["statusDetails"]
        )
        ):
            document.add_heading("Details", level=2)
            if "message" in test["statusDetails"]:
                document.add_paragraph(test["statusDetails"]["message"], style=None)
            if "trace" in config_info and "trace" in test["statusDetails"]:
                table = document.add_table(rows=1, cols=1, style="Trace table")
                hdr_cells = table.rows[0].cells
                hdr_cells[0].add_paragraph(test["statusDetails"]["trace"] + "\n", style="Code")

        if "links" in config_info and "links" in test and len(test["links"]) != 0:
            document.add_heading("Links", level=2)
            for link in test["links"]:
                if "name" in link and "url" in link:
                    document.add_paragraph("{}: {}".format(link["name"], link["url"]))
                else:
                    print("WARNING: A link was provided without name or url and will not be printed.")

        if "setup" in config_info:
            document.add_heading("Test Setup", level=2)
            for parent in test["parents"]:
                if "befores" in parent:
                    for before in parent["befores"]:
                        document.add_paragraph("[Fixture] {}".format(before["name"]), style="Step")
                        print_attachments(document, before)
                        print_steps(document, before, config_info, 1)
            if document.paragraphs[-1].text == "Test Setup":
                delete_paragraph(document.paragraphs[-1])

        if "body" in config_info:
            document.add_heading("Test Body", level=2)
            print_attachments(document, test)
            print_steps(document, test, config_info)
            if document.paragraphs[-1].text == "Test Body":
                delete_paragraph(document.paragraphs[-1])

        if "teardown" in config_info:
            document.add_heading("Test Teardown", level=2)
            for parent in test["parents"]:
                if "afters" in parent:
                    for after in parent["afters"]:
                        document.add_paragraph("[Fixture] {}".format(after["name"]), style="Step")
                        print_attachments(document, after)
                        print_steps(document, after, config_info, 1)
            if document.paragraphs[-1].text == "Test Teardown":
                delete_paragraph(document.paragraphs[-1])

    document.save(config['output_filename'])


def run(
        alluredir,
        template_path,
        output_filename,
        title,
        logo_path,
        logo_height,
        config_option,
):
    results, session = build_data(alluredir)

    imgfile = os.path.join(session["alluredir"], "pie.png")
    session["piechart_source"] = imgfile
    piechart.create_piechart(session["results"], imgfile)
    config_path = config_option
    script_path = os.path.dirname(os.path.realpath(__file__))
    standard_config_path = script_path + "/config/standard.ini"
    if config_option == "standard":
        config_path = standard_config_path
    if config_option == "standard_on_fail":
        config_path = script_path + "/config/standard_on_fail.ini"
    elif config_option == "compact":
        config_path = script_path + "/config/compact.ini"
    elif config_option == "no_trace":
        config_path = script_path + "/config/no_trace.ini"
    config = get_config_from_file(standard_config_path, config_path)
    config['logo'] = {}
    config['logo']['path'] = logo_path
    config['logo']['height'] = logo_height
    config['output_filename'] = output_filename
    config['template_path'] = template_path
    if 'title' not in config['cover']:
        config['cover']['title'] = title

    create_docx(results, session, config)
