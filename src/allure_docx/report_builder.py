import logging
import os
import re
import warnings
import shutil
import subprocess
import json
import matplotlib.pyplot as plt

from os import listdir
from os.path import join, isfile
from time import ctime
from datetime import timedelta, datetime

from docx.shared import Mm, Cm
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx2pdf import convert


class ReportBuilder:
    """
    Builder to create a report from a given ReportConfig Object.
    """

    def __init__(self, allure_dir, config):
        self.indent = 6
        self.config = config
        self.config['allure_dir'] = allure_dir
        if 'template_path' not in self.config:
            self.config['template_path'] = os.path.join(os.path.dirname(os.path.realpath(__file__)), "template.docx")
        self.document = Document(config['template_path'])

        self.session = {
            "allure_dir": config['allure_dir'],
            "start": None,
            "stop": None,
            "results": {
                "passed": 0,
                "skipped": 0,
                "broken": 0,
                "failed": 0,
                "unknown": 0
            },
            "results_relative": {
                "passed": 0,
                "skipped": 0,
                "broken": 0,
                "failed": 0,
                "unknown": 0
            },
            "total": 0,
        }

        self.sorted_recent_results = None
        try:
            self._build_data()
            self._create_pie_chart()
            self._print_report()
        except KeyError:
            raise KeyError(f"A key error for the json data occurred. Your json data may be corrupt or possibly an "
                           "unsupported framework was used that. See https://github.com/typhoon-hil/allure-docx/issues"
                           " if this is a known issue or create a new one if it is not.")

    def save_report(self, output):
        """
        Save report to given output path as docx.
        """
        self.document.save(output)

    def save_report_to_pdf(self, output):
        """
        Save report to given output path as pdf. Tries officetopdf or soffice.
        """

        soffice = shutil.which("soffice")

        temp_docx_filename = f"{os.path.dirname(output)}/__temp.docx"
        temp_pdf_filename = f"{os.path.dirname(output)}/__temp.pdf"
        self.save_report(temp_docx_filename)

        try:
            convert(temp_docx_filename, output)
        except Exception as e:  # noqa
            if soffice is not None:
                libre_version_string = subprocess.check_output(["soffice", "--version"]).decode("utf-8")
                libre_version_match = re.search(r"(\d+)\.\d+\.\d+\.\d+", libre_version_string)
                libre_version = libre_version_match.group()
                libre_major_version = libre_version_match.group(1)
                if int(libre_major_version) < 7:
                    logging.warning("Working with Libre Office version " + libre_version
                                    + " to generate pdf from docx. Version > 7 is recommended. Bugs like faulty color "
                                      "or missing line breaks may appear.")
                result_dir = os.path.dirname(output)
                subprocess.call(["soffice", "--convert-to", "pdf", "--outdir", result_dir, temp_docx_filename])
                os.rename(temp_pdf_filename, output)
            else:
                print("Failed to convert via docx2pdf or soffice.")
                print(str(e))

        os.remove(temp_docx_filename)

    def _process_steps(self, node):
        """
        Check starting and stopping time of each step and adjust start time in session dict accordingly.
        """
        if "start" in node:
            if "start" not in self.session:
                self.session["start"] = node["start"]
            elif self.session["start"] is None:
                self.session["start"] = node["start"]
            elif node["start"] < self.session["start"]:
                self.session["start"] = node["start"]

        if "stop" in node:
            if "stop" not in self.session:
                self.session["stop"] = node["stop"]
            elif self.session["stop"] is None:
                self.session["stop"] = node["stop"]
            elif node["stop"] > self.session["stop"]:
                self.session["stop"] = node["stop"]

        if "steps" in node:
            for step in node["steps"]:
                self._process_steps(step)

    def _build_data(self):
        """
        Build the session dict and the sorted_results dict from the given allure directory.
        """

        def get_sorting_key(d):
            classification = {"broken": 0, "failed": 1, "skipped": 2, "passed": 3}
            return f"{classification[d['status']]}-{d['name']}"

        allure_dir = self.config['allure_dir']

        json_results = [f for f in listdir(allure_dir) if isfile(join(allure_dir, f)) and "result" in f]
        json_containers = [f for f in listdir(allure_dir) if isfile(join(allure_dir, f)) and "container" in f]

        data_containers = []
        for file_name in json_containers:
            with open(join(allure_dir, file_name), encoding="utf-8") as file:
                container = json.load(file)
                data_containers.append(container)

        data_results_dict = {}
        for file_name in json_results:  # one array of results per test historyId
            with open(join(allure_dir, file_name), encoding="utf-8") as file:
                result = json.load(file)
                history_id = result['historyId']
                if history_id not in data_results_dict:
                    data_results_dict[history_id] = []
                data_results_dict[history_id].append(result)
        history_data_results = list(data_results_dict.items())  # can be used in a later version to implement history
        for tests in history_data_results:
            tests[1].sort(key=lambda x: x["start"], reverse=True)
        recent_results = [results[1][0] for results in history_data_results]  # get only the most recent results
        id_sorted_recent_results = sorted(recent_results, key=lambda x: x["fullName"])

        idx = -1
        param_idx = 1
        for result in id_sorted_recent_results:
            idx += 1
            # create unique names for parameterized tests
            if "parameters" in result and len(result["parameters"]) > 0:
                if idx > 0 and result["fullName"] == id_sorted_recent_results[idx - 1]["fullName"]:
                    result["name"] += f" [{param_idx}]"
                    if param_idx == 1:
                        id_sorted_recent_results[idx - 1]["name"] += " [0]"
                    param_idx += 1
                else:
                    param_idx = 1

            self.sorted_recent_results = sorted(id_sorted_recent_results, key=get_sorting_key)

            self._process_steps(result)
            self.session["total"] += 1
            self.session["results"][result["status"]] += 1

            result["parents"] = []
            for container in data_containers:
                if "children" not in container:
                    continue
                if result["uuid"] not in container["children"]:
                    continue
                result["parents"].append(container)
                if "befores" in container:
                    for before in container["befores"]:
                        self._process_steps(before)
                if "afters" in container:
                    for after in container["afters"]:
                        self._process_steps(after)

        if self.session["total"] == 0:
            warnings.warn("No test result files were found!")

        if self.session["start"] is not None:
            self.session["duration"] = self.session["stop"] - self.session["start"]
            self.session["start"] = ctime(self.session["start"] / 1000.0)
            self.session["stop"] = ctime(self.session["stop"] / 1000.0)
        else:
            self.session["duration"] = "Not available"
            self.session["start"] = "Not available"
            self.session["stop"] = "Not available"

        for item in self.session["results"]:
            if self.session["total"] > 0:
                self.session["results_relative"][item] = "{:.2f}%".format(
                    100 * self.session["results"][item] / self.session["total"])
            else:
                self.session["results_relative"][item] = "Not available"

    def _create_pie_chart(self):
        """
        Creates the pie chart for allure results overview and saves it into the allure_dir folder.
        """
        img_file = os.path.join(self.session["allure_dir"], "pie.png")
        self.session["pie_chart_source"] = img_file

        color_map = {
            "passed": "#97CC64",
            "broken": "#FFD050",
            "failed": "#FD5A3E",
            "skipped": "#AAAAAA",
            "unknown": "#D35EBE"
        }
        colors = []
        data_arr = []
        labels = []
        for item in self.session["results"]:
            if self.session["results"][item] != 0:
                data_arr.append(self.session["results"][item])
                colors.append(color_map[item])
                labels.append(item)

        fig, ax = plt.subplots()
        ax.pie(data_arr, startangle=90, wedgeprops=dict(width=0.5), labels=data_arr, labeldistance=0.7, colors=colors)
        ax.legend(labels, frameon=False, loc='upper left', bbox_to_anchor=(-0.25, 0, 0.5, 1))
        fig.savefig(img_file, bbox_inches="tight")

    def _print_report(self):
        """
        Main function to print the docx document. Raises Error if no allure result files were found.
        """
        if not self.sorted_recent_results:
            raise ImportError("No test result files were found in the given allure results folder.")

        self._print_cover()
        self.document.add_section()

        footer = self.document.sections[1].footer
        footer.is_linked_to_previous = False
        self._print_footer(footer)
        header = self.document.sections[1].header
        header.is_linked_to_previous = False
        self._print_header(header, True)

        self._print_details()
        self._print_session_summary()

        # print tests
        for test in self.sorted_recent_results:
            # print only the most recent test, history could be included later.
            if "tests" in self.config["info"][test["status"]]:
                self._print_test(test)

    def _print_attachments(self, item):
        """
        Print attachments from allure results to the document.
        """
        if "attachments" in item:
            for attachment in item["attachments"]:
                if 'name' not in attachment:
                    attachment['name'] = ""
                self.document.add_paragraph(f"[Attachment] {attachment['name']}", style="Step")
                if "image" in attachment["type"]:
                    self.document.add_picture(
                        os.path.join(self.session["allure_dir"], attachment["source"]),
                        width=Mm(100),
                    )
                    self.document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.LEFT

    @staticmethod
    def _format_argval(argval):
        """Remove newlines and limit max length

        From Allure-pytest logger (formats argument in the CLI live logs).
        Consider using the same function."""

        max_arg_length = 100
        argval = argval.replace("\n", " ")
        if len(argval) > max_arg_length:
            argval = argval[:3] + " ... " + argval[-max_arg_length:]
        return argval

    def _print_steps(self, parent_step, config_info, indent=0):
        """
        Print the given step. The info sub dict of the given test must be provided to apply the configuration.
        If there are sub-steps in the given parent-step this methods recursively calls itself with the given step.
        For each Recursion the indent for the print is incremented.
        """
        indent_str = indent * self.indent * " "
        if "steps" in parent_step:
            for step in parent_step["steps"]:
                if step["status"] in ["failed", "broken"]:
                    step_style = "Step Failed"
                else:
                    step_style = "Step"
                self.document.add_paragraph(f"{indent_str}> {step['name']}", style=step_style)
                if "parameters" in config_info and "parameters" in step:
                    for params in step["parameters"]:
                        paragraph = self.document.add_paragraph(f"{indent_str}    ", style="Step Param Parag")
                        paragraph.add_run(
                            f"{params['name']} = {self._format_argval(params['value'])}",
                            style="Step Param",
                        )
                if "details" in config_info and "statusDetails" in step and len(step["statusDetails"]) != 0:
                    if "message" in step["statusDetails"] and len(step["statusDetails"]["message"]) != 0:
                        self.document.add_paragraph(step["statusDetails"]["message"], style=step_style)

                    if "trace" in config_info and "trace" in step["statusDetails"] and len(
                            step["statusDetails"]["trace"]) != 0:
                        table = self.document.add_table(rows=1, cols=1, style="Trace table")
                        hdr_cells = table.rows[0].cells
                        hdr_cells[0].add_paragraph(step["statusDetails"]["trace"] + "\n", style="Code")
                        self.document.add_paragraph("", style=None)
                self._print_steps(step, config_info, indent + 1)
                if "attachments" in config_info:
                    self._print_attachments(step)

    @staticmethod
    def _add_field(run, field):
        """
        Creates a docx field and appends it to the given run object.
        """

        def create_attribute(element, name, value):
            element.set(qn(name), value)

        def create_element(name):
            return OxmlElement(name)

        fld_char1 = create_element('w:fldChar')
        create_attribute(fld_char1, 'w:fldCharType', 'begin')

        instr_text = create_element('w:instrText')
        create_attribute(instr_text, 'xml:space', 'preserve')
        instr_text.text = field

        fld_char2 = create_element('w:fldChar')
        create_attribute(fld_char2, 'w:fldCharType', 'end')

        run._r.append(fld_char1)
        run._r.append(instr_text)
        run._r.append(fld_char2)

    def _print_footer(self, footer):
        """
        Prints a footer to the given footer object, including date and page number.
        """
        footer.paragraphs[0].text += datetime.today().strftime('%Y-%m-%d')
        footer.paragraphs[0].text += "\t\t"
        footer_run = footer.paragraphs[0].add_run()
        self._add_field(footer_run, field="PAGE")

    @staticmethod
    def _delete_paragraph(paragraph):
        """
        Deletes a given paragraph from the document.
        """
        p_element = paragraph._element
        p_element.getparent().remove(p_element)
        p_element._p = p_element._element = None

    def _print_header(self, header, details=False):
        """
        Prints a header to the given header object. This includes a logo (if a logo is specified)
        and test details if details set to True. Details include the title and the "Device under test" if specified.
        """
        htable = header.add_table(1, 2, Cm(16))
        htable.style = "header table"
        htab_cells = htable.rows[0].cells

        if 'logo' in self.config:
            ht1 = htab_cells[1].add_paragraph()
            ht1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            kh = ht1.add_run()
            if 'width' in self.config['logo']:
                logo_width = Cm(self.config['logo'])
            else:
                logo_width = Cm(5)
            kh.add_picture(self.config['logo']['path'], width=logo_width)

        if details:
            header_text = self.config['cover']['title']
            if 'Device under test' in self.config['details']:
                header_text += "\n" + self.config['details']['Device under test']
            htab_cells[0].add_paragraph(header_text)
            header.add_paragraph("")

        self._delete_paragraph(header.paragraphs[0])

        # creates corrupted files:
        # self._delete_paragraph(htab_cells[0].paragraphs[0])
        # self._delete_paragraph(htab_cells[1].paragraphs[0])

    def _print_cover(self):
        """
        Prints the cover page to the document. This includes title, date and if specified inside the [cover] section,
        the company name and device under test.
        """
        header = self.document.sections[0].header
        self._print_header(header)

        self._delete_paragraph(self.document.paragraphs[0])
        if 'company' in self.config['cover']:
            self.document.add_paragraph("\n" + self.config['cover']['company'], style="company")
        self.document.add_paragraph("\n\n\n\nTest Report", style="Title")
        subtitle = self.config['cover']['title']
        if 'Device under test' in self.config['details']:
            subtitle += "\n" + self.config['details']['Device under test']
        self.document.add_paragraph(subtitle, style="Subtitle")
        self.document.add_paragraph("\n" + datetime.today().strftime('%Y-%m-%d'), style="heading 2")

    def _print_details(self):
        """
        Prints the test details that are specified inside the [details] section of the configuration file
        and a table of content for the tests.
        """

        if 'details' in self.config and len(self.config['details']) > 0:
            self.document.add_paragraph("Test Details", style="Heading 1")
            i = 0
            detail_table = self.document.add_table(rows=len(self.config['details']), cols=2, style="Label table")
            for detail in self.config['details'].items():
                detail_table.rows[i].cells[0].paragraphs[-1].clear().add_run(detail[0])
                detail_table.rows[i].cells[1].paragraphs[-1].clear().add_run(detail[1].strip())
                i += 1

            detail_table.columns[0].width = Cm(4)
            for cell in detail_table.columns[0].cells:
                cell.width = Cm(4)
            detail_table.columns[1].width = Cm(12)
            for cell in detail_table.columns[1].cells:
                cell.width = Cm(12)
            self.document.add_page_break()

    def _print_session_summary(self):
        """
        Prints the session summary, including results, total running time and a pie chart.
        """

        has_session_summary = (int(self.config["summary"]["overview"]) is not 0
                               or int(self.config["summary"]["table"]) is not 0)

        if has_session_summary:
            self.document.add_paragraph("Test Session Summary", style="Heading 1")

        if int(self.config["summary"]["overview"]) != 0:
            table = self.document.add_table(rows=1, cols=2)
            summary_cell = table.rows[0].cells[0]
            duration_string = self._duration_to_string(self.session["duration"])
            summary_cell.add_paragraph(
                f"Start: {self.session['start']}\nEnd: {self.session['stop']}\nDuration: {duration_string}"
            )
            self._delete_paragraph(summary_cell.paragraphs[0])

            results_strs = []
            for item in self.session["results"]:
                results_strs.append(
                    f"{item}: {self.session['results'][item]} ({self.session['results_relative'][item]})")
            summary_cell.add_paragraph("\n".join(results_strs))

            pie_chart_cell = table.rows[0].cells[1]
            paragraph = pie_chart_cell.paragraphs[0]
            run = paragraph.add_run()
            run.add_picture(self.session["pie_chart_source"], width=Mm(75))

            self.document.add_paragraph("")

        results = self.session['results']

        def print_result_table(status):
            if results[status] > 0:
                result_table = self.document.add_table(rows=results[status], cols=2, style=f"{status} table")
                i = 0
                for test in self.sorted_recent_results:
                    if test['status'] == status:
                        result_table.rows[i].cells[0].paragraphs[-1].add_run(
                            test['name'])
                        result_table.rows[i].cells[1].paragraphs[-1].add_run(
                            status)
                        i += 1

                result_table.columns[0].width = Cm(12)
                for cell in result_table.columns[0].cells:
                    cell.width = Cm(12)
                result_table.columns[1].width = Cm(4)
                for cell in result_table.columns[1].cells:
                    cell.width = Cm(4)

        if int(self.config["summary"]["table"]) != 0:
            if "tests" in self.config["info"]["failed"]:
                print_result_table("failed")
            if "tests" in self.config["info"]["broken"]:
                print_result_table("broken")
            if "tests" in self.config["info"]["skipped"]:
                print_result_table("skipped")
            if "tests" in self.config["info"]["passed"]:
                print_result_table("passed")

        if has_session_summary:
            self.document.add_page_break()

    @staticmethod
    def _duration_to_string(ms) -> str:
        """
        Parses milliseconds to a fitting string with unit.
        """
        duration_unit = "ms"
        if ms < 1000:
            return str(ms) + "ms"
        seconds = ms / 1000.0
        if seconds < 60:
            return str(int(seconds)) + "s"
        minutes = seconds / 60.0
        if minutes < 60:
            return str(int(minutes)) + "m " + str(int(seconds % 60)) + "s"
        hours = minutes / 60
        return str(int(hours)) + "h " + str(int(minutes % 60 % 60)) + "m " + str(int(seconds % 60)) + "s"

    def _print_test(self, test):
        """
        Prints the specified test to the document.
        """
        # config elements for the specific status of this test
        config_info = self.config["info"][test["status"]]
        config_labels = self.config["labels"][test["status"]]

        self.document.add_paragraph(f"{test['name']}  [ {test['status']} ]", style=f"Heading {test['status']}")

        table = None
        added_table = False
        if "duration" in config_info:
            duration = test["stop"] - test["start"]
            duration_string = self._duration_to_string(duration)

            table = self.document.add_table(rows=1, cols=2, style="Label table")
            table.rows[0].cells[0].paragraphs[-1].clear().add_run("Duration")
            table.rows[0].cells[1].paragraphs[-1].clear().add_run(duration_string)
            added_table = True

        # add labels to table
        for label_name in config_labels:
            if not added_table:
                table = self.document.add_table(rows=0, cols=2, style="Label table")
                added_table = True
            iterator = iter(label for label in test["labels"] if label["name"].lower() == label_name)
            label = next(iterator, None)
            if label is not None:
                row = table.add_row()
                row.cells[0].paragraphs[-1].clear().add_run(label_name.capitalize())
                while label is not None:
                    row.cells[1].add_paragraph(label["value"])
                    label = next(iterator, None)
                self._delete_paragraph(row.cells[1].paragraphs[0])

        if table is not None:
            table.columns[0].width = Cm(4)
            for cell in table.columns[0].cells:
                cell.width = Cm(4)
            table.columns[1].width = Cm(12)
            for cell in table.columns[1].cells:
                cell.width = Cm(12)
            self.document.add_paragraph()

        if "description" in config_info:
            self.document.add_heading("Description", level=2)
            if "description" in test and len(test["description"]) != 0:
                self.document.add_paragraph(test["description"])

        if "parameters" in config_info and "parameters" in test and len(test["parameters"]) != 0:
            self.document.add_heading("Parameters", level=2)
            for p in test["parameters"]:
                self.document.add_paragraph(f"{p['name']}: {p['value']}", style="Step")

        if (
                "details" in config_info
                and "statusDetails" in test
                and len(test["statusDetails"]) != 0
                and (
                "message" in test["statusDetails"]
                and len(test["statusDetails"]["message"]) != 0
                or "trace" in config_info
                and "trace" in test["statusDetails"]
        )
        ):
            self.document.add_heading("Details", level=2)
            if "message" in test["statusDetails"]:
                self.document.add_paragraph(test["statusDetails"]["message"], style=None)
            if "trace" in config_info and "trace" in test["statusDetails"]:
                table = self.document.add_table(rows=1, cols=1, style="Trace table")
                hdr_cells = table.rows[0].cells
                hdr_cells[0].add_paragraph(test["statusDetails"]["trace"] + "\n", style="Code")
                self.document.add_paragraph("", style=None)

        if "links" in config_info and "links" in test and len(test["links"]) != 0:
            self.document.add_heading("Links", level=2)
            for link in test["links"]:
                if "name" in link and "url" in link:
                    self.document.add_paragraph(f"{link['name']}: {link['url']}")
                else:
                    print("WARNING: A link was provided without name or url and will not be printed.")

        if "setup" in config_info:
            self.document.add_heading("Test Setup", level=2)
            for parent in test["parents"]:
                if "befores" in parent:
                    for before in parent["befores"]:
                        self.document.add_paragraph(f"[Fixture] {before['name']}", style="Step")
                        if "steps" in config_info:
                            self._print_steps(before, config_info, 1)
                        self._print_attachments(before)
            if self.document.paragraphs[-1].text == "Test Setup":
                self._delete_paragraph(self.document.paragraphs[-1])

        if "body" in config_info:
            self.document.add_heading("Test Body", level=2)
            if "steps" in config_info:
                self._print_steps(test, config_info)
            self._print_attachments(test)
            if self.document.paragraphs[-1].text == "Test Body":
                self._delete_paragraph(self.document.paragraphs[-1])

        if "teardown" in config_info:
            self.document.add_heading("Test Teardown", level=2)
            for parent in test["parents"]:
                if "afters" in parent:
                    for after in parent["afters"]:
                        self.document.add_paragraph(f"[Fixture] {after['name']}", style="Step")
                        if "steps" in config_info:
                            self._print_steps(after, config_info, 1)
                        self._print_attachments(after)
            if self.document.paragraphs[-1].text == "Test Teardown":
                self._delete_paragraph(self.document.paragraphs[-1])

        self.document.add_paragraph("", style=None)
